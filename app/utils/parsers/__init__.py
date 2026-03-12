"""
文档解析工具模块
支持 PDF、Word、Markdown、纯文本等多种格式
"""
from pathlib import Path
from typing import Optional, Dict, Any
import asyncio

from app.core.logger import logger


class DocumentParser:
    """文档解析器基类"""

    async def parse(self, file_path: str) -> str:
        """
        解析文档，提取文本内容

        Args:
            file_path: 文件路径

        Returns:
            提取的文本内容
        """
        raise NotImplementedError


class TextParser(DocumentParser):
    """纯文本解析器"""

    async def parse(self, file_path: str) -> str:
        """解析纯文本文件"""
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                content = f.read()
            logger.info(f"✅ 文本文件解析成功: {file_path}")
            return content
        except UnicodeDecodeError:
            # 尝试其他编码
            with open(file_path, "r", encoding="gbk") as f:
                content = f.read()
            logger.info(f"✅ 文本文件解析成功 (GBK): {file_path}")
            return content


class MarkdownParser(DocumentParser):
    """Markdown 解析器"""

    async def parse(self, file_path: str) -> str:
        """解析 Markdown 文件"""
        with open(file_path, "r", encoding="utf-8") as f:
            content = f.read()
        logger.info(f"✅ Markdown 文件解析成功: {file_path}")
        return content


class PDFParser(DocumentParser):
    """PDF 解析器"""

    async def parse(self, file_path: str) -> str:
        """解析 PDF 文件"""
        try:
            import pypdf

            text_parts = []
            with open(file_path, "rb") as f:
                pdf_reader = pypdf.PdfReader(f)
                num_pages = len(pdf_reader.pages)

                for page_num in range(num_pages):
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)

            content = "\n\n".join(text_parts)
            logger.info(f"✅ PDF 文件解析成功 ({num_pages}页): {file_path}")
            return content

        except Exception as e:
            logger.error(f"❌ PDF 解析失败: {e}")
            raise


class WordParser(DocumentParser):
    """Word 文档解析器"""

    async def parse(self, file_path: str) -> str:
        """解析 Word 文档 (.docx)"""
        try:
            from docx import Document

            doc = Document(file_path)
            text_parts = []

            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)

            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells])
                    if row_text.strip():
                        text_parts.append(row_text)

            content = "\n\n".join(text_parts)
            logger.info(f"✅ Word 文档解析成功: {file_path}")
            return content

        except Exception as e:
            logger.error(f"❌ Word 解析失败: {e}")
            raise


def get_parser(file_type: str) -> DocumentParser:
    """
    根据文件类型获取对应的解析器

    Args:
        file_type: 文件类型 (pdf/txt/md/docx)

    Returns:
        对应的解析器实例

    Raises:
        ValueError: 不支持的文件类型
    """
    parsers = {
        "txt": TextParser(),
        "md": MarkdownParser(),
        "pdf": PDFParser(),
        "docx": WordParser(),
    }

    if file_type not in parsers:
        raise ValueError(f"不支持的文件类型: {file_type}")

    return parsers[file_type]


async def parse_document(file_path: str, file_type: str) -> str:
    """
    解析文档（自动选择解析器）

    Args:
        file_path: 文件路径
        file_type: 文件类型

    Returns:
        提取的文本内容
    """
    parser = get_parser(file_type)
    return await parser.parse(file_path)


__all__ = [
    "DocumentParser",
    "TextParser",
    "MarkdownParser",
    "PDFParser",
    "WordParser",
    "get_parser",
    "parse_document",
]
